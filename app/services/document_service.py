"""
Document processing and management service
"""

import os
import uuid
import hashlib
import json
import re
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import asyncio
from io import BytesIO

from sqlalchemy.orm import Session
from sqlalchemy import desc, and_, or_
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import redis

from config import settings
from app.models.document import Document, DocumentStatus, DocumentAccess
from app.models.template import Template, Placeholder
from app.models.user import User
from app.schemas.document import (
    DocumentCreate, DocumentUpdate, DocumentGenerate, DocumentShare,
    DocumentSearch, DocumentStats, DocumentPreview
)
from app.services.encryption_service import EncryptionService
from database import get_db

# Redis client for caching
redis_client = redis.Redis(
    host=settings.REDIS_HOST,
    port=settings.REDIS_PORT,
    decode_responses=True
)


class DocumentService:
    """Document processing and management service"""

    @staticmethod
    def create_document(db: Session, document_data: DocumentCreate, user_id: int) -> Document:
        """Create a new document"""

        # Generate file path if template is used
        file_path = None
        if document_data.template_id:
            filename = f"{uuid.uuid4()}.{document_data.file_format}"
            file_path = os.path.join(settings.DOCUMENTS_PATH, filename)

        document = Document(
            title=document_data.title,
            description=document_data.description,
            content=document_data.content,
            placeholder_data=document_data.placeholder_data,
            file_path=file_path,
            file_format=document_data.file_format,
            access_level=document_data.access_level,
            requires_signature=document_data.requires_signature,
            required_signature_count=document_data.required_signature_count,
            auto_delete=document_data.auto_delete,
            user_id=user_id,
            template_id=document_data.template_id,
            status=DocumentStatus.DRAFT
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        return document

    @staticmethod
    def create_document_from_generation(db: Session, generation_data: DocumentGenerate, user_id: int) -> Document:
        """Create document from generation request"""

        filename = f"{uuid.uuid4()}.{generation_data.file_format}"
        file_path = os.path.join(settings.DOCUMENTS_PATH, filename)

        document = Document(
            title=generation_data.title,
            description=generation_data.description,
            placeholder_data=generation_data.placeholder_data,
            file_path=file_path,
            file_format=generation_data.file_format,
            access_level=generation_data.access_level,
            requires_signature=generation_data.requires_signature,
            required_signature_count=generation_data.required_signature_count,
            user_id=user_id,
            template_id=generation_data.template_id,
            status=DocumentStatus.PROCESSING
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        return document

    @staticmethod
    def create_document_for_batch(db: Session, template_id: int, doc_data: Dict[str, Any],
                                user_id: int, batch_id: str) -> Document:
        """Create document for batch processing"""

        filename = f"{uuid.uuid4()}.docx"
        file_path = os.path.join(settings.DOCUMENTS_PATH, filename)

        document = Document(
            title=doc_data["title"],
            description=doc_data.get("description"),
            placeholder_data=doc_data["placeholder_data"],
            file_path=file_path,
            file_format="docx",
            user_id=user_id,
            template_id=template_id,
            status=DocumentStatus.PROCESSING
        )

        # Add batch metadata
        document.placeholder_data["_batch_id"] = batch_id

        db.add(document)
        db.commit()
        db.refresh(document)

        return document

    @staticmethod
    def update_document(db: Session, document: Document, document_update: DocumentUpdate) -> Document:
        """Update document"""

        for field, value in document_update.dict(exclude_unset=True).items():
            setattr(document, field, value)

        document.updated_at = datetime.utcnow()
        db.commit()
        db.refresh(document)

        return document

    @staticmethod
    def delete_document(db: Session, document: Document) -> None:
        """Soft delete document"""

        document.status = DocumentStatus.ARCHIVED
        document.deleted_at = datetime.utcnow()
        db.commit()

        # Remove file if exists
        if document.file_path and os.path.exists(document.file_path):
            try:
                os.remove(document.file_path)
            except OSError:
                pass

    @staticmethod
    def search_documents(db: Session, user_id: int, search_params: DocumentSearch) -> Tuple[List[Document], int]:
        """Search documents with advanced filters"""

        query = db.query(Document).filter(Document.user_id == user_id)

        # Apply filters
        if search_params.query:
            query = query.filter(
                or_(
                    Document.title.contains(search_params.query),
                    Document.description.contains(search_params.query),
                    Document.content.contains(search_params.query)
                )
            )

        if search_params.status:
            query = query.filter(Document.status == search_params.status)

        if search_params.template_id:
            query = query.filter(Document.template_id == search_params.template_id)

        if search_params.created_after:
            query = query.filter(Document.created_at >= search_params.created_after)

        if search_params.created_before:
            query = query.filter(Document.created_at <= search_params.created_before)

        # Get total count
        total = query.count()

        # Apply sorting
        if search_params.sort_by == "created_at":
            order_col = Document.created_at
        elif search_params.sort_by == "title":
            order_col = Document.title
        elif search_params.sort_by == "status":
            order_col = Document.status
        else:
            order_col = Document.updated_at

        if search_params.sort_order == "asc":
            query = query.order_by(order_col)
        else:
            query = query.order_by(desc(order_col))

        # Apply pagination
        documents = query.offset(
            (search_params.page - 1) * search_params.per_page
        ).limit(search_params.per_page).all()

        return documents, total

    @staticmethod
    def get_user_document_stats(db: Session, user_id: int) -> DocumentStats:
        """Get document statistics for user"""

        # Total documents
        total_documents = db.query(Document).filter(Document.user_id == user_id).count()

        # By status
        completed_documents = db.query(Document).filter(
            Document.user_id == user_id,
            Document.status == DocumentStatus.COMPLETED
        ).count()

        draft_documents = db.query(Document).filter(
            Document.user_id == user_id,
            Document.status == DocumentStatus.DRAFT
        ).count()

        failed_documents = db.query(Document).filter(
            Document.user_id == user_id,
            Document.status == DocumentStatus.FAILED
        ).count()

        # Size calculation
        documents_with_size = db.query(Document).filter(
            Document.user_id == user_id,
            Document.file_size.isnot(None)
        ).all()
        total_size = sum(doc.file_size or 0 for doc in documents_with_size)

        # Monthly stats
        current_month_start = datetime.utcnow().replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        last_month_start = (current_month_start - timedelta(days=1)).replace(day=1)

        this_month = db.query(Document).filter(
            Document.user_id == user_id,
            Document.created_at >= current_month_start
        ).count()

        last_month = db.query(Document).filter(
            Document.user_id == user_id,
            Document.created_at >= last_month_start,
            Document.created_at < current_month_start
        ).count()

        # Calculate growth rate
        if last_month > 0:
            growth_rate = ((this_month - last_month) / last_month) * 100
        else:
            growth_rate = 100.0 if this_month > 0 else 0.0

        return DocumentStats(
            total_documents=total_documents,
            completed_documents=completed_documents,
            draft_documents=draft_documents,
            failed_documents=failed_documents,
            total_size=total_size,
            this_month=this_month,
            last_month=last_month,
            growth_rate=growth_rate
        )

    @staticmethod
    async def generate_document_from_template(document_id: int, placeholder_data: Dict[str, Any]) -> bool:
        """Generate document from template with placeholders"""

        try:
            db = next(get_db())

            # Get document and template
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                return False

            template = db.query(Template).filter(Template.id == document.template_id).first()
            if not template:
                document.status = DocumentStatus.FAILED
                document.error_message = "Template not found"
                db.commit()
                return False

            start_time = datetime.utcnow()

            # Load template document
            template_path = os.path.join(settings.TEMPLATES_PATH, template.file_path)
            if not os.path.exists(template_path):
                document.status = DocumentStatus.FAILED
                document.error_message = "Template file not found"
                db.commit()
                return False

            # Generate document
            success = await DocumentService._process_template_placeholders(
                template_path, document.file_path, template, placeholder_data
            )

            if success:
                # Update document status
                document.status = DocumentStatus.COMPLETED
                document.completed_at = datetime.utcnow()
                document.generation_time = (datetime.utcnow() - start_time).total_seconds()

                # Calculate file size and hash
                if os.path.exists(document.file_path):
                    document.file_size = os.path.getsize(document.file_path)
                    document.file_hash = DocumentService._calculate_file_hash(document.file_path)

                # Encrypt if required
                if document.is_encrypted:
                    encrypted_path = await EncryptionService.encrypt_file(document.file_path)
                    if encrypted_path:
                        document.file_path = encrypted_path
                        document.encryption_key_id = "default"  # Use default encryption key
            else:
                document.status = DocumentStatus.FAILED
                document.error_message = "Failed to process template"

            db.commit()
            db.close()

            return success

        except Exception as e:
            # Update document with error
            try:
                document.status = DocumentStatus.FAILED
                document.error_message = str(e)
                db.commit()
            except:
                pass
            finally:
                db.close()

            return False

    @staticmethod
    async def _process_template_placeholders(template_path: str, output_path: str,
                                           template: Template, placeholder_data: Dict[str, Any]) -> bool:
        """Process template placeholders and generate output document"""

        try:
            # Load template document
            doc = DocxDocument(template_path)

            # Get placeholders for this template
            db = next(get_db())
            placeholders = db.query(Placeholder).filter(
                Placeholder.template_id == template.id
            ).order_by(
                Placeholder.paragraph_index,
                Placeholder.start_run_index
            ).all()
            db.close()

            # Process placeholders
            for placeholder in placeholders:
                value = placeholder_data.get(placeholder.name, placeholder.default_value or "")

                # Format value based on placeholder type
                formatted_value = DocumentService._format_placeholder_value(
                    value, placeholder.placeholder_type, placeholder.casing
                )

                # Apply formatting
                DocumentService._replace_placeholder_in_document(
                    doc, placeholder, formatted_value
                )

            # Apply document-level formatting
            DocumentService._apply_document_formatting(doc, template)

            # Save document
            doc.save(output_path)

            return True

        except Exception as e:
            print(f"Error processing template: {e}")
            return False

    @staticmethod
    def _format_placeholder_value(value: str, placeholder_type: str, casing: str) -> str:
        """Format placeholder value based on type and casing"""

        # Handle different types
        if placeholder_type == "date" and value:
            try:
                from dateutil.parser import parse
                date_obj = parse(value)
                return date_obj.strftime("%B %d, %Y")
            except:
                pass

        # Apply casing
        if casing == "upper":
            return value.upper()
        elif casing == "lower":
            return value.lower()
        elif casing == "title":
            return value.title()

        return value

    @staticmethod
    def _replace_placeholder_in_document(doc: DocxDocument, placeholder: Placeholder, value: str):
        """Replace placeholder in document with formatted value"""

        try:
            paragraph = doc.paragraphs[placeholder.paragraph_index]

            # Find the placeholder text in runs
            placeholder_text = f"${{{placeholder.name}}}"

            # Simple replacement - in production, you'd want more sophisticated handling
            full_text = "".join(run.text for run in paragraph.runs)
            if placeholder_text in full_text:
                # Clear all runs and add the replaced text
                for run in paragraph.runs:
                    run.clear()

                # Add new text with formatting
                new_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                new_run.text = full_text.replace(placeholder_text, value)

                # Apply formatting
                new_run.bold = placeholder.bold
                new_run.italic = placeholder.italic
                new_run.underline = placeholder.underline

        except (IndexError, AttributeError):
            # Skip if paragraph index is invalid
            pass

    @staticmethod
    def _apply_document_formatting(doc: DocxDocument, template: Template):
        """Apply document-level formatting"""

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = template.font_family
        font.size = Pt(template.font_size)

        # Apply margins if specified
        if template.page_margins:
            sections = doc.sections
            for section in sections:
                margins = template.page_margins
                if margins.get('top'):
                    section.top_margin = Inches(margins['top'])
                if margins.get('bottom'):
                    section.bottom_margin = Inches(margins['bottom'])
                if margins.get('left'):
                    section.left_margin = Inches(margins['left'])
                if margins.get('right'):
                    section.right_margin = Inches(margins['right'])

    @staticmethod
    def _calculate_file_hash(file_path: str) -> str:
        """Calculate SHA256 hash of file"""

        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    @staticmethod
    def create_share_link(db: Session, document: Document, share_data: DocumentShare) -> Dict[str, Any]:
        """Create shareable link for document"""

        # Generate share token
        share_token = str(uuid.uuid4())

        # Set expiration
        expires_at = None
        if share_data.expires_in_days:
            expires_at = datetime.utcnow() + timedelta(days=share_data.expires_in_days)

        # Update document
        document.share_token = share_token
        document.share_expires_at = expires_at
        document.access_level = share_data.access_level

        db.commit()

        # Cache share info in Redis
        share_info = {
            "document_id": document.id,
            "access_level": share_data.access_level.value,
            "password_protected": share_data.password_protected,
            "expires_at": expires_at.isoformat() if expires_at else None
        }

        if share_data.password_protected and share_data.password:
            # Hash password for storage
            from passlib.context import CryptContext
            pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
            share_info["password_hash"] = pwd_context.hash(share_data.password)

        redis_client.setex(
            f"share:{share_token}",
            86400 * 30,  # 30 days
            json.dumps(share_info)
        )

        return {
            "share_token": share_token,
            "share_url": f"/api/documents/shared/{share_token}",
            "expires_at": expires_at,
            "access_level": share_data.access_level.value
        }

    @staticmethod
    def validate_shared_access(document: Document, password: Optional[str] = None) -> Dict[str, Any]:
        """Validate access to shared document"""

        # Check if share token exists
        if not document.share_token:
            return {"valid": False, "error": "Document is not shared"}

        # Check expiration
        if document.share_expires_at and document.share_expires_at < datetime.utcnow():
            return {"valid": False, "error": "Share link has expired"}

        # Get share info from Redis
        share_info_str = redis_client.get(f"share:{document.share_token}")
        if not share_info_str:
            return {"valid": False, "error": "Share link is invalid"}

        try:
            share_info = json.loads(share_info_str)
        except json.JSONDecodeError:
            return {"valid": False, "error": "Invalid share data"}

        # Check password if required
        if share_info.get("password_protected") and share_info.get("password_hash"):
            if not password:
                return {"valid": False, "error": "Password required"}

            from passlib.context import CryptContext
            pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

            if not pwd_context.verify(password, share_info["password_hash"]):
                return {"valid": False, "error": "Invalid password"}

        return {"valid": True}

    @staticmethod
    def generate_preview(document: Document) -> DocumentPreview:
        """Generate document preview"""

        # Create preview content
        preview_content = ""
        if document.content:
            # Truncate content for preview
            preview_content = document.content[:500] + "..." if len(document.content) > 500 else document.content
        elif document.placeholder_data:
            # Create preview from placeholder data
            preview_lines = []
            for key, value in document.placeholder_data.items():
                if not key.startswith("_"):  # Skip internal keys
                    preview_lines.append(f"{key}: {value}")
            preview_content = "\n".join(preview_lines[:10])  # First 10 lines

        # Generate real thumbnail
        thumbnail_result = None
        page_count = 1

        if document.file_path and os.path.exists(document.file_path):
            try:
                from app.services.thumbnail_service import thumbnail_service
                # Use synchronous thumbnail generation for preview
                thumbnail_result = {"success": False, "thumbnail_url": None}
                # TODO: Implement sync thumbnail generation or skip for now
            except ImportError:
                thumbnail_result = {"success": False, "thumbnail_url": None}

            # Calculate actual page count for PDFs
            if document.file_format.lower() == 'pdf':
                page_count = DocumentService._get_pdf_page_count(document.file_path)

        return DocumentPreview(
            id=document.id,
            title=document.title,
            content_preview=preview_content,
            thumbnail_url=thumbnail_result.get('thumbnail_url') if thumbnail_result and thumbnail_result['success'] else None,
            page_count=page_count,
            watermarked=True
        )

    @staticmethod
    def cleanup_expired_documents():
        """Cleanup expired documents (background task)"""

        db = next(get_db())

        try:
            # Find documents that should be auto-deleted
            expired_documents = db.query(Document).filter(
                Document.auto_delete == True,
                Document.retention_expires_at < datetime.utcnow()
            ).all()

            for document in expired_documents:
                # Remove file
                if document.file_path and os.path.exists(document.file_path):
                    try:
                        os.remove(document.file_path)
                    except OSError:
                        pass

                # Delete from database
                db.delete(document)

            db.commit()

        finally:
            db.close()

    @staticmethod
    def get_processing_queue_status() -> Dict[str, Any]:
        """Get status of document processing queue"""
        try:
            from celery import current_app
            import redis
            
            # Get Celery app instance
            celery_app = current_app
            
            # Get Redis connection for queue inspection
            redis_client = redis.from_url(settings.CELERY_BROKER_URL)
            
            # Get queue statistics from Redis
            queue_name = 'document_processing'
            pending_count = redis_client.llen(queue_name)
            
            # Get active tasks from Celery
            inspect = celery_app.control.inspect()
            active_tasks = inspect.active()
            processing_count = sum(len(tasks) for tasks in active_tasks.values()) if active_tasks else 0
            
            # Get reserved tasks (queued but not yet active)
            reserved_tasks = inspect.reserved()
            reserved_count = sum(len(tasks) for tasks in reserved_tasks.values()) if reserved_tasks else 0
            
            # Get failed tasks from recent history
            stats = inspect.stats()
            failed_count = 0
            if stats:
                for worker_stats in stats.values():
                    failed_count += worker_stats.get('failed', 0)
            
            # Calculate queue length (pending + reserved)
            queue_length = pending_count + reserved_count
            
            # Get processing time statistics from Redis cache
            avg_processing_time = 2.5  # Default fallback
            try:
                cached_avg = redis_client.get("document_processing:avg_time")
                if cached_avg:
                    avg_processing_time = float(cached_avg)
            except:
                pass
            
            return {
                "pending_documents": pending_count,
                "processing_documents": processing_count,
                "failed_documents": failed_count,
                "queue_length": queue_length,
                "average_processing_time": avg_processing_time,
                "workers_online": len(active_tasks) if active_tasks else 0
            }
            
        except Exception as e:
            logger.error(f"Failed to get queue status: {e}")
            # Fallback to basic status
            return {
                "pending_documents": 0,
                "processing_documents": 0,
                "failed_documents": 0,
                "queue_length": 0,
                "average_processing_time": 2.5,
                "workers_online": 0,
                "error": "Queue monitoring unavailable"
            }

    @staticmethod
    async def consolidate_batch_inputs(
        db: Session,
        templates: List[Template],
        batch_documents: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """Smart input consolidation for batch processing"""

        from app.services.batch_processing_service import BatchInputConsolidator

        return await BatchInputConsolidator.consolidate_batch_inputs(
            db, templates, batch_documents
        )

    @staticmethod
    async def create_batch_download(
        documents: List[Document],
        batch_id: str
    ) -> Dict[str, Any]:
        """Create ZIP file for batch document download"""

        import zipfile
        import tempfile

        try:
            # Create temporary ZIP file
            temp_dir = tempfile.mkdtemp()
            zip_path = os.path.join(temp_dir, f"batch_{batch_id}.zip")

            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for doc in documents:
                    if doc.file_path and os.path.exists(doc.file_path):
                        # Add file to ZIP with proper naming
                        filename = f"{doc.title}.{doc.file_format}"
                        # Sanitize filename
                        filename = re.sub(r'[^\w\s.-]', '', filename)
                        zipf.write(doc.file_path, filename)

            # Get file size
            file_size = os.path.getsize(zip_path)

            return {
                "success": True,
                "file_path": zip_path,
                "file_size": file_size
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }

    @staticmethod
    async def generate_document_preview(document_id: int) -> Dict[str, Any]:
        """Generate preview for a document"""
        try:
            import os
            from PIL import Image
            import fitz  # PyMuPDF
            from pdf2image import convert_from_path
            
            # Get document from database
            db = next(get_db())
            try:
                document = db.query(Document).filter(Document.id == document_id).first()
                if not document or not os.path.exists(document.file_path):
                    return {"error": "Document not found"}
                
                file_path = document.file_path
                file_size = os.path.getsize(file_path)
                file_size_mb = round(file_size / (1024 * 1024), 2)
                
                # Create preview directory
                preview_dir = os.path.join(settings.STORAGE_PATH, "previews")
                os.makedirs(preview_dir, exist_ok=True)
                
                preview_filename = f"preview_{document_id}.png"
                thumbnail_filename = f"thumb_{document_id}.png"
                preview_path = os.path.join(preview_dir, preview_filename)
                thumbnail_path = os.path.join(preview_dir, thumbnail_filename)
                
                page_count = 1
                
                # Generate preview based on file type
                if document.file_format.lower() == 'pdf':
                    try:
                        # Use PyMuPDF for better performance
                        pdf_doc = fitz.open(file_path)
                        page_count = pdf_doc.page_count
                        
                        # Get first page as image
                        first_page = pdf_doc[0]
                        pix = first_page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # 2x zoom for quality
                        
                        # Save preview
                        pix.save(preview_path)
                        
                        # Create thumbnail
                        img = Image.open(preview_path)
                        img.thumbnail((300, 300), Image.Resampling.LANCZOS)
                        img.save(thumbnail_path)
                        
                        pdf_doc.close()
                        
                    except Exception as e:
                        logger.warning(f"PyMuPDF failed, trying pdf2image: {e}")
                        try:
                            # Fallback to pdf2image
                            images = convert_from_path(file_path, first_page=1, last_page=1, dpi=150)
                            if images:
                                images[0].save(preview_path, 'PNG')
                                
                                # Create thumbnail
                                img = images[0].copy()
                                img.thumbnail((300, 300), Image.Resampling.LANCZOS)
                                img.save(thumbnail_path)
                                
                        except Exception as e2:
                            logger.error(f"Preview generation failed: {e2}")
                            return {"error": "Could not generate preview"}
                
                elif document.file_format.lower() in ['docx', 'doc']:
                    try:
                        # Convert DOCX to PDF first, then to image
                        from docx2pdf import convert
                        temp_pdf = os.path.join(preview_dir, f"temp_{document_id}.pdf")
                        
                        # Convert to PDF
                        convert(file_path, temp_pdf)
                        
                        # Generate preview from PDF
                        pdf_doc = fitz.open(temp_pdf)
                        page_count = pdf_doc.page_count
                        
                        first_page = pdf_doc[0]
                        pix = first_page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                        pix.save(preview_path)
                        
                        # Create thumbnail
                        img = Image.open(preview_path)
                        img.thumbnail((300, 300), Image.Resampling.LANCZOS)
                        img.save(thumbnail_path)
                        
                        pdf_doc.close()
                        
                        # Clean up temp PDF
                        os.remove(temp_pdf)
                        
                    except Exception as e:
                        logger.error(f"DOCX preview generation failed: {e}")
                        # Create placeholder image
                        placeholder = Image.new('RGB', (600, 800), color='white')
                        placeholder.save(preview_path)
                        placeholder.thumbnail((300, 300))
                        placeholder.save(thumbnail_path)
                
                elif document.file_format.lower() in ['png', 'jpg', 'jpeg']:
                    try:
                        # For image files, just create thumbnail
                        import shutil
                        shutil.copy(file_path, preview_path)
                        
                        img = Image.open(file_path)
                        img.thumbnail((300, 300), Image.Resampling.LANCZOS)
                        img.save(thumbnail_path)
                        
                    except Exception as e:
                        logger.error(f"Image preview generation failed: {e}")
                        return {"error": "Could not process image"}
                
                else:
                    # Unsupported format - create placeholder
                    placeholder = Image.new('RGB', (600, 800), color='#f0f0f0')
                    placeholder.save(preview_path)
                    placeholder.thumbnail((300, 300))
                    placeholder.save(thumbnail_path)
                
                return {
                    "preview_url": f"/api/documents/{document_id}/preview",
                    "thumbnail_url": f"/api/documents/{document_id}/thumbnail",
                    "file_size": f"{file_size_mb} MB",
                    "page_count": page_count,
                    "generated_at": datetime.utcnow().isoformat()
                }
                
            finally:
                db.close()

        except Exception as e:
            logger.error(f"Preview generation failed for document {document_id}: {e}")
            return {"error": str(e)}

    @staticmethod
    def _get_pdf_page_count(file_path: str) -> int:
        """Calculate actual page count for PDF files"""

        try:
            # Try PyMuPDF first
            try:
                import fitz
                pdf_doc = fitz.open(file_path)
                page_count = pdf_doc.page_count
                pdf_doc.close()
                return page_count
            except ImportError:
                pass

            # Fallback to basic PDF parsing
            with open(file_path, 'rb') as f:
                content = f.read()
                # Count /Type /Page occurrences
                import re
                pages = re.findall(rb'/Type\s*/Page[^s]', content)
                return len(pages) if pages else 1

        except Exception as e:
            logger.warning(f"Could not determine PDF page count: {e}")
            return 1
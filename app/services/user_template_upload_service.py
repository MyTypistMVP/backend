"""
User Template Upload Service
Allow users to upload documents, extract placeholders, and make them public/private
"""

import json
import logging
import re
import os
import uuid
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from sqlalchemy.orm import Session
from sqlalchemy import Column, Integer, String, Text, DateTime, ForeignKey, Boolean, Float, desc
from database import Base
from docx import Document
import docx2txt

logger = logging.getLogger(__name__)


class UserUploadedTemplate(Base):
    """User-uploaded templates with extracted placeholders"""
    __tablename__ = "user_uploaded_templates"

    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=False, index=True)

    # Template details
    title = Column(String(255), nullable=False)
    description = Column(Text, nullable=True)
    category = Column(String(100), nullable=True)
    tags = Column(Text, nullable=True)  # JSON array of tags

    # File information
    original_filename = Column(String(255), nullable=False)
    file_path = Column(String(500), nullable=False)
    file_size_bytes = Column(Integer, nullable=False)
    file_type = Column(String(50), nullable=False)  # docx, pdf, etc.

    # Extraction results
    extracted_placeholders = Column(Text, nullable=True)  # JSON array of placeholders
    placeholder_count = Column(Integer, default=0)
    extraction_status = Column(String(20), default="pending")  # pending, completed, failed
    extraction_error = Column(Text, nullable=True)

    # Visibility and approval
    visibility = Column(String(20), default="private")  # private, public, pending_review
    admin_approved = Column(Boolean, default=False)
    admin_review_notes = Column(Text, nullable=True)
    reviewed_by = Column(Integer, ForeignKey('users.id'), nullable=True)
    reviewed_at = Column(DateTime, nullable=True)

    # Pricing (set by admin for public templates)
    price_tokens = Column(Integer, nullable=True)
    admin_set_price = Column(Boolean, default=False)

    # Usage statistics
    views_count = Column(Integer, default=0)
    downloads_count = Column(Integer, default=0)
    revenue_generated = Column(Float, default=0.0)

    # Status
    is_active = Column(Boolean, default=True, index=True)
    created_at = Column(DateTime, default=datetime.utcnow, index=True)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class PlaceholderExtractionLog(Base):
    """Log of placeholder extraction attempts"""
    __tablename__ = "placeholder_extraction_logs"

    id = Column(Integer, primary_key=True, index=True)
    template_id = Column(Integer, ForeignKey('user_uploaded_templates.id'), nullable=False, index=True)

    # Extraction details
    extraction_method = Column(String(50), nullable=False)  # regex, ml, manual
    placeholders_found = Column(Integer, default=0)
    extraction_time_ms = Column(Integer, nullable=True)

    # Results
    success = Column(Boolean, default=False)
    error_message = Column(Text, nullable=True)
    extracted_data = Column(Text, nullable=True)  # JSON of extraction details

    created_at = Column(DateTime, default=datetime.utcnow)


class UserTemplateUploadService:
    """Service for handling user template uploads and placeholder extraction"""

    # Common placeholder patterns
    PLACEHOLDER_PATTERNS = [
        r'\{([^}]+)\}',  # {placeholder}
        r'\[\[([^\]]+)\]\]',  # [[placeholder]]
        r'\{\{([^}]+)\}\}',  # {{placeholder}}
        r'_([A-Z_]+)_',  # _PLACEHOLDER_
        r'<([^>]+)>',  # <placeholder>
        r'\$\{([^}]+)\}',  # ${placeholder}
    ]

    # Common placeholder types based on name
    PLACEHOLDER_TYPES = {
        'name': 'text',
        'first_name': 'text',
        'last_name': 'text',
        'full_name': 'text',
        'email': 'email',
        'phone': 'phone',
        'address': 'address',
        'date': 'date',
        'signature': 'signature',
        'company': 'text',
        'title': 'text',
        'amount': 'number',
        'price': 'number',
        'quantity': 'number'
    }

    @staticmethod
    def upload_user_template(
        db: Session,
        user_id: int,
        file_data: bytes,
        filename: str,
        title: str,
        description: Optional[str] = None,
        category: Optional[str] = None,
        tags: Optional[List[str]] = None,
        visibility: str = "private"
    ) -> Dict[str, Any]:
        """Upload and process user template"""
        try:
            # Validate file
            if not filename.lower().endswith(('.docx', '.doc')):
                return {
                    "success": False,
                    "error": "Only Word documents (.docx, .doc) are supported"
                }

            # Generate unique file path
            file_extension = os.path.splitext(filename)[1]
            unique_filename = f"{uuid.uuid4()}{file_extension}"
            upload_dir = "uploads/user_templates"
            os.makedirs(upload_dir, exist_ok=True)
            file_path = os.path.join(upload_dir, unique_filename)

            # Save file
            with open(file_path, 'wb') as f:
                f.write(file_data)

            # Create template record
            template = UserUploadedTemplate(
                user_id=user_id,
                title=title,
                description=description,
                category=category,
                tags=json.dumps(tags or []),
                original_filename=filename,
                file_path=file_path,
                file_size_bytes=len(file_data),
                file_type=file_extension[1:],  # Remove dot
                visibility=visibility,
                extraction_status="pending"
            )

            db.add(template)
            db.commit()
            db.refresh(template)

            # Extract placeholders asynchronously (or immediately for demo)
            extraction_result = UserTemplateUploadService._extract_placeholders(
                db, template.id, file_path
            )

            logger.info(f"User template uploaded: {template.id} by user {user_id}")

            return {
                "success": True,
                "template_id": template.id,
                "title": title,
                "file_path": file_path,
                "extraction_status": extraction_result.get("status", "pending"),
                "placeholders_found": extraction_result.get("placeholders_count", 0),
                "message": "Template uploaded successfully. Placeholder extraction in progress."
            }

        except Exception as e:
            logger.error(f"Failed to upload user template: {e}")
            raise

    @staticmethod
    def _extract_placeholders(
        db: Session,
        template_id: int,
        file_path: str
    ) -> Dict[str, Any]:
        """Extract placeholders from uploaded document"""
        start_time = datetime.utcnow()

        try:
            # Get template record
            template = db.query(UserUploadedTemplate).filter(
                UserUploadedTemplate.id == template_id
            ).first()

            if not template:
                return {"success": False, "error": "Template not found"}

            # Extract text from document
            document_text = UserTemplateUploadService._extract_text_from_document(file_path)

            # Find placeholders using multiple methods
            placeholders = UserTemplateUploadService._find_placeholders_multi_method(
                document_text, file_path
            )

            # Process and categorize placeholders
            processed_placeholders = UserTemplateUploadService._process_placeholders(placeholders)

            # Update template record
            template.extracted_placeholders = json.dumps(processed_placeholders)
            template.placeholder_count = len(processed_placeholders)
            template.extraction_status = "completed"
            template.updated_at = datetime.utcnow()

            # Log extraction
            extraction_time = int((datetime.utcnow() - start_time).total_seconds() * 1000)

            extraction_log = PlaceholderExtractionLog(
                template_id=template_id,
                extraction_method="multi_method",
                placeholders_found=len(processed_placeholders),
                extraction_time_ms=extraction_time,
                success=True,
                extracted_data=json.dumps({
                    "raw_placeholders": placeholders,
                    "processed_placeholders": processed_placeholders,
                    "document_length": len(document_text)
                })
            )

            db.add(extraction_log)
            db.commit()

            logger.info(f"Placeholders extracted for template {template_id}: {len(processed_placeholders)} found")

            return {
                "success": True,
                "status": "completed",
                "placeholders_count": len(processed_placeholders),
                "placeholders": processed_placeholders,
                "extraction_time_ms": extraction_time
            }

        except Exception as e:
            # Log error
            extraction_log = PlaceholderExtractionLog(
                template_id=template_id,
                extraction_method="multi_method",
                placeholders_found=0,
                success=False,
                error_message=str(e)
            )

            db.add(extraction_log)

            # Update template status
            template = db.query(UserUploadedTemplate).filter(
                UserUploadedTemplate.id == template_id
            ).first()

            if template:
                template.extraction_status = "failed"
                template.extraction_error = str(e)
                template.updated_at = datetime.utcnow()

            db.commit()

            logger.error(f"Failed to extract placeholders for template {template_id}: {e}")

            return {
                "success": False,
                "status": "failed",
                "error": str(e)
            }

    @staticmethod
    def _extract_text_from_document(file_path: str) -> str:
        """Extract text from Word document"""
        try:
            # Try python-docx first
            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

            document_text = '\n'.join(text_parts)

            # Fallback to docx2txt if needed
            if not document_text.strip():
                document_text = docx2txt.process(file_path)

            return document_text

        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            # Fallback method
            try:
                return docx2txt.process(file_path)
            except:
                return ""

    @staticmethod
    def _find_placeholders_multi_method(document_text: str, file_path: str) -> List[Dict[str, Any]]:
        """Find placeholders using multiple extraction methods"""
        all_placeholders = []

        # Method 1: Regex patterns
        for pattern in UserTemplateUploadService.PLACEHOLDER_PATTERNS:
            matches = re.finditer(pattern, document_text, re.IGNORECASE)
            for match in matches:
                placeholder_name = match.group(1).strip()
                if placeholder_name and len(placeholder_name) > 1:
                    all_placeholders.append({
                        "name": placeholder_name,
                        "pattern": pattern,
                        "method": "regex",
                        "position": match.start(),
                        "original_text": match.group(0)
                    })

        # Method 2: Look for common words that might be placeholders
        placeholder_keywords = [
            'name', 'date', 'address', 'phone', 'email', 'signature',
            'company', 'title', 'amount', 'price', 'quantity', 'description'
        ]

        for keyword in placeholder_keywords:
            # Look for patterns like "Name: ______" or "Date: _____"
            pattern = rf'{keyword}\s*:?\s*[_\s]{{3,}}'
            matches = re.finditer(pattern, document_text, re.IGNORECASE)
            for match in matches:
                all_placeholders.append({
                    "name": keyword,
                    "pattern": "keyword_detection",
                    "method": "keyword",
                    "position": match.start(),
                    "original_text": match.group(0)
                })

        # Method 3: Look for bracketed content that might be instructions
        instruction_pattern = r'\[([^\]]{10,})\]'  # Longer bracketed content
        matches = re.finditer(instruction_pattern, document_text)
        for match in matches:
            content = match.group(1).strip()
            # Check if it looks like a placeholder instruction
            if any(word in content.lower() for word in ['enter', 'insert', 'fill', 'type', 'write']):
                # Try to extract the field name
                field_name = UserTemplateUploadService._extract_field_name_from_instruction(content)
                if field_name:
                    all_placeholders.append({
                        "name": field_name,
                        "pattern": "instruction",
                        "method": "instruction",
                        "position": match.start(),
                        "original_text": match.group(0),
                        "instruction": content
                    })

        return all_placeholders

    @staticmethod
    def _extract_field_name_from_instruction(instruction: str) -> Optional[str]:
        """Extract field name from instruction text"""
        instruction_lower = instruction.lower()

        # Common patterns in instructions
        patterns = [
            r'enter\s+(?:your\s+)?(\w+)',
            r'insert\s+(?:your\s+)?(\w+)',
            r'fill\s+(?:in\s+)?(?:your\s+)?(\w+)',
            r'type\s+(?:your\s+)?(\w+)',
            r'write\s+(?:your\s+)?(\w+)',
        ]

        for pattern in patterns:
            match = re.search(pattern, instruction_lower)
            if match:
                return match.group(1)

        # Fallback: look for key words
        key_words = ['name', 'date', 'address', 'phone', 'email', 'signature', 'company']
        for word in key_words:
            if word in instruction_lower:
                return word

        return None

    @staticmethod
    def _process_placeholders(raw_placeholders: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Process and deduplicate placeholders"""
        # Deduplicate by name (case-insensitive)
        unique_placeholders = {}

        for placeholder in raw_placeholders:
            name = placeholder["name"].lower().strip()

            # Clean up name
            name = re.sub(r'[^a-zA-Z0-9_]', '_', name)
            name = re.sub(r'_+', '_', name).strip('_')

            if name and len(name) > 1:
                if name not in unique_placeholders:
                    # Determine placeholder type
                    placeholder_type = UserTemplateUploadService._determine_placeholder_type(name)

                    unique_placeholders[name] = {
                        "name": name,
                        "display_name": UserTemplateUploadService._generate_display_name(name),
                        "type": placeholder_type,
                        "required": True,
                        "validation": UserTemplateUploadService._get_validation_rules(placeholder_type),
                        "extraction_method": placeholder["method"],
                        "original_text": placeholder.get("original_text", ""),
                        "instruction": placeholder.get("instruction"),
                        "position": placeholder.get("position", 0)
                    }

        # Convert to list and sort by position
        processed_placeholders = list(unique_placeholders.values())
        processed_placeholders.sort(key=lambda x: x["position"])

        return processed_placeholders

    @staticmethod
    def _determine_placeholder_type(name: str) -> str:
        """Determine placeholder type based on name"""
        name_lower = name.lower()

        for key, ptype in UserTemplateUploadService.PLACEHOLDER_TYPES.items():
            if key in name_lower:
                return ptype

        # Additional pattern matching
        if 'mail' in name_lower:
            return 'email'
        elif 'phone' in name_lower or 'tel' in name_lower:
            return 'phone'
        elif 'date' in name_lower or 'time' in name_lower:
            return 'date'
        elif 'sign' in name_lower:
            return 'signature'
        elif 'address' in name_lower or 'location' in name_lower:
            return 'address'
        elif 'amount' in name_lower or 'price' in name_lower or 'cost' in name_lower:
            return 'number'
        else:
            return 'text'

    @staticmethod
    def _generate_display_name(name: str) -> str:
        """Generate human-readable display name"""
        # Convert snake_case to Title Case
        display_name = name.replace('_', ' ').title()

        # Handle common abbreviations
        replacements = {
            'Id': 'ID',
            'Url': 'URL',
            'Api': 'API',
            'Dob': 'Date of Birth',
            'Ssn': 'SSN'
        }

        for old, new in replacements.items():
            display_name = display_name.replace(old, new)

        return display_name

    @staticmethod
    def _get_validation_rules(placeholder_type: str) -> Dict[str, Any]:
        """Get validation rules for placeholder type"""
        validation_rules = {
            'text': {
                'min_length': 1,
                'max_length': 255,
                'pattern': None
            },
            'email': {
                'pattern': r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$',
                'error_message': 'Please enter a valid email address'
            },
            'phone': {
                'pattern': r'^\+?[\d\s\-\(\)]{10,}$',
                'error_message': 'Please enter a valid phone number'
            },
            'number': {
                'type': 'number',
                'min': 0
            },
            'date': {
                'type': 'date',
                'format': 'YYYY-MM-DD'
            },
            'signature': {
                'type': 'file',
                'accept': 'image/*',
                'max_size': 5 * 1024 * 1024  # 5MB
            },
            'address': {
                'type': 'textarea',
                'max_length': 500
            }
        }

        return validation_rules.get(placeholder_type, validation_rules['text'])

    @staticmethod
    def get_user_templates(
        db: Session,
        user_id: int,
        visibility: Optional[str] = None,
        limit: int = 20
    ) -> Dict[str, Any]:
        """Get user's uploaded templates"""
        try:
            query = db.query(UserUploadedTemplate).filter(
                UserUploadedTemplate.user_id == user_id,
                UserUploadedTemplate.is_active == True
            )

            if visibility:
                query = query.filter(UserUploadedTemplate.visibility == visibility)

            templates = query.order_by(desc(UserUploadedTemplate.created_at)).limit(limit).all()

            templates_data = []
            for template in templates:
                # Parse placeholders
                placeholders = []
                if template.extracted_placeholders:
                    try:
                        placeholders = json.loads(template.extracted_placeholders)
                    except:
                        pass

                # Parse tags
                tags = []
                if template.tags:
                    try:
                        tags = json.loads(template.tags)
                    except:
                        pass

                templates_data.append({
                    "template_id": template.id,
                    "title": template.title,
                    "description": template.description,
                    "category": template.category,
                    "tags": tags,
                    "original_filename": template.original_filename,
                    "file_size_bytes": template.file_size_bytes,
                    "placeholder_count": template.placeholder_count,
                    "placeholders": placeholders,
                    "visibility": template.visibility,
                    "admin_approved": template.admin_approved,
                    "price_tokens": template.price_tokens,
                    "views_count": template.views_count,
                    "downloads_count": template.downloads_count,
                    "extraction_status": template.extraction_status,
                    "created_at": template.created_at.isoformat(),
                    "updated_at": template.updated_at.isoformat()
                })

            return {
                "success": True,
                "templates": templates_data,
                "total_count": len(templates_data)
            }

        except Exception as e:
            logger.error(f"Failed to get user templates: {e}")
            raise

    @staticmethod
    def update_template_visibility(
        db: Session,
        template_id: int,
        user_id: int,
        visibility: str
    ) -> Dict[str, Any]:
        """Update template visibility"""
        try:
            template = db.query(UserUploadedTemplate).filter(
                UserUploadedTemplate.id == template_id,
                UserUploadedTemplate.user_id == user_id
            ).first()

            if not template:
                return {"success": False, "error": "Template not found"}

            # Validate visibility
            valid_visibilities = ["private", "public", "pending_review"]
            if visibility not in valid_visibilities:
                return {"success": False, "error": "Invalid visibility setting"}

            # If making public, set to pending_review for admin approval
            if visibility == "public" and not template.admin_approved:
                visibility = "pending_review"

            template.visibility = visibility
            template.updated_at = datetime.utcnow()

            db.commit()

            return {
                "success": True,
                "template_id": template_id,
                "visibility": visibility,
                "message": "Template visibility updated successfully"
            }

        except Exception as e:
            logger.error(f"Failed to update template visibility: {e}")
            raise

    @staticmethod
    def get_admin_review_queue(
        db: Session,
        limit: int = 50
    ) -> Dict[str, Any]:
        """Get templates pending admin review"""
        try:
            templates = db.query(UserUploadedTemplate).filter(
                UserUploadedTemplate.visibility == "pending_review",
                UserUploadedTemplate.is_active == True
            ).order_by(UserUploadedTemplate.created_at).limit(limit).all()

            templates_data = []
            for template in templates:
                # Get user info
                from app.models.user import User
                user = db.query(User).filter(User.id == template.user_id).first()

                # Parse placeholders
                placeholders = []
                if template.extracted_placeholders:
                    try:
                        placeholders = json.loads(template.extracted_placeholders)
                    except:
                        pass

                templates_data.append({
                    "template_id": template.id,
                    "title": template.title,
                    "description": template.description,
                    "category": template.category,
                    "user": {
                        "id": user.id if user else None,
                        "name": f"{user.first_name} {user.last_name}" if user else "Unknown",
                        "email": user.email if user else None
                    },
                    "original_filename": template.original_filename,
                    "placeholder_count": template.placeholder_count,
                    "placeholders": placeholders,
                    "extraction_status": template.extraction_status,
                    "created_at": template.created_at.isoformat(),
                    "file_path": template.file_path
                })

            return {
                "success": True,
                "templates": templates_data,
                "total_count": len(templates_data)
            }

        except Exception as e:
            logger.error(f"Failed to get admin review queue: {e}")
            raise

    @staticmethod
    def admin_approve_template(
        db: Session,
        template_id: int,
        admin_user_id: int,
        approved: bool,
        price_tokens: Optional[int] = None,
        review_notes: Optional[str] = None
    ) -> Dict[str, Any]:
        """Admin approve or reject user template"""
        try:
            template = db.query(UserUploadedTemplate).filter(
                UserUploadedTemplate.id == template_id
            ).first()

            if not template:
                return {"success": False, "error": "Template not found"}

            template.admin_approved = approved
            template.reviewed_by = admin_user_id
            template.reviewed_at = datetime.utcnow()
            template.admin_review_notes = review_notes

            if approved:
                template.visibility = "public"
                if price_tokens is not None:
                    template.price_tokens = price_tokens
                    template.admin_set_price = True
            else:
                template.visibility = "private"

            template.updated_at = datetime.utcnow()

            db.commit()

            # Send notification to user (would integrate with notification service)
            logger.info(f"Template {template_id} {'approved' if approved else 'rejected'} by admin {admin_user_id}")

            return {
                "success": True,
                "template_id": template_id,
                "approved": approved,
                "visibility": template.visibility,
                "price_tokens": template.price_tokens,
                "message": f"Template {'approved' if approved else 'rejected'} successfully"
            }

        except Exception as e:
            logger.error(f"Failed to admin approve template: {e}")
            raise


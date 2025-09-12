"""Guest document generation and customization utilities"""

from typing import Dict, Any, Optional
from pathlib import Path
import json
from fastapi import HTTPException
import asyncio
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

from app.models.template import Template
from app.models.document import Document, DocumentStatus
from app.utils.storage import get_storage_path, ensure_storage_path
from app.services.document_service import DocumentService
from app.utils.security import sanitize_content


async def generate_guest_preview(
    template: Template,
    customization: Dict[str, Any],
    session_id: str
) -> str:
    """Generate a preview document for guest users"""
    
    # Sanitize and validate customization data
    safe_data = sanitize_content(customization)
    
    # Create preview document path
    preview_dir = get_storage_path() / "previews" / session_id
    ensure_storage_path(preview_dir)
    preview_path = preview_dir / f"preview_{template.id}.docx"
    
    # Load template
    template_doc = DocxDocument(template.file_path)
    
    # Apply customizations with limits for guest users
    await apply_guest_customizations(template_doc, safe_data)
    
    # Save preview
    template_doc.save(str(preview_path))
    
    return str(preview_path)


async def apply_guest_customizations(
    doc: DocxDocument,
    customization: Dict[str, Any]
) -> None:
    """Apply limited customizations for guest preview"""
    
    # Limit customization options for guests
    allowed_fonts = ["Arial", "Times New Roman", "Calibri"]
    max_font_size = 14
    
    for paragraph in doc.paragraphs:
        # Apply text replacements
        for key, value in customization.get("replacements", {}).items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
        
        # Apply limited formatting
        font = customization.get("font")
        if font and font in allowed_fonts:
            for run in paragraph.runs:
                run.font.name = font
        
        size = customization.get("size")
        if size and 8 <= size <= max_font_size:
            for run in paragraph.runs:
                run.font.size = Pt(size)


async def finalize_guest_document(
    doc: Document,
    customization: Dict[str, Any],
    session_id: str
) -> Document:
    """Finalize a guest document with applied customizations"""
    
    # Ensure document is in guest state
    if doc.status != DocumentStatus.GUEST:
        raise HTTPException(
            status_code=400,
            detail="Can only finalize guest documents"
        )
    
    # Apply customizations
    doc_path = Path(doc.file_path)
    if doc_path.exists():
        template_doc = DocxDocument(doc_path)
        await apply_guest_customizations(template_doc, customization)
        
        # Save as new version
        final_path = doc_path.parent / f"final_{doc.id}.docx"
        template_doc.save(str(final_path))
        doc.file_path = str(final_path)
    
    # Update metadata
    doc.metadata = {
        **(doc.metadata or {}),
        "guest_customization": customization,
        "guest_session": session_id
    }
    
    return doc